"""
file_reader.py - 多格式文件读取
路径：scripts/file_reader.py
版本：v2.3.5-part2-hotfix1.1 - 版本统一(Claude Code 系统修复)
"""
import os, sys, json, hashlib, chardet
from pathlib import Path
from datetime import datetime

PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))


class FileReader:
    SUPPORTED_TYPES = {
        ".pdf":"pdf", ".docx":"word", ".xlsx":"excel", ".xls":"excel_legacy",
        ".csv":"csv", ".txt":"text", ".md":"text",
        ".jpg":"image", ".jpeg":"image", ".png":"image",
    }

    def __init__(self, config=None):
        if config is None:
            p = PROJECT_ROOT / "config" / "settings.json"
            config = json.load(open(p,"r",encoding="utf-8")) if p.exists() else {}
        self.config = config
        self.allowed_paths = config.get("allowed_paths", [str(PROJECT_ROOT/"data")])
        self.max_file_size_mb = config.get("max_file_size_mb", 50)

    def check_file_access(self, fp):
        ap = os.path.abspath(fp)
        for a in self.allowed_paths:
            if ap.startswith(os.path.abspath(a)): return True
        raise PermissionError(f"安全拦截: 文件不在允许路径内: {fp}")

    def get_file_hash(self, fp):
        h = hashlib.md5()
        with open(fp,"rb") as f:
            for chunk in iter(lambda: f.read(8192), b""): h.update(chunk)
        return h.hexdigest()

    def read_file(self, file_path):
        result = {"file_path":str(file_path),"file_type":"","file_size_mb":0,"file_hash":"",
                  "content":"","metadata":{},"pages":0,"read_method":"","success":False,"error":""}
        try:
            self.check_file_access(file_path)
            size_mb = os.path.getsize(file_path)/(1024*1024)
            if size_mb > self.max_file_size_mb:
                result["error"] = f"文件过大: {size_mb:.1f}MB > {self.max_file_size_mb}MB"; return result
            result["file_size_mb"] = size_mb
            result["file_hash"] = self.get_file_hash(file_path)
            ext = Path(file_path).suffix.lower()
            ft = self.SUPPORTED_TYPES.get(ext)
            if not ft: result["error"] = f"不支持的格式: {ext}"; return result
            result["file_type"] = ft

            if ft == "pdf": self._read_pdf(file_path, result)
            elif ft == "word": self._read_docx(file_path, result)
            elif ft == "excel": self._read_xlsx(file_path, result)
            elif ft == "csv": self._read_csv(file_path, result)
            elif ft == "text": self._read_text(file_path, result)
            elif ft == "image": self._read_image(file_path, result)
            elif ft in ("excel_legacy",): result["error"] = "请转换为.xlsx后重试"; return result

            result["success"] = True
        except Exception as e:
            result["error"] = f"{type(e).__name__}: {e}"
        return result

    def _read_pdf(self, fp, r):
        import pdfplumber
        parts = []
        with pdfplumber.open(fp) as pdf:
            r["pages"] = len(pdf.pages)
            r["metadata"] = {"page_count": len(pdf.pages)}
            for i, page in enumerate(pdf.pages):
                t = page.extract_text()
                if t and t.strip(): parts.append(f"[第{i+1}页]\n{t.strip()}")
                for ti, table in enumerate(page.extract_tables()):
                    tt = self._fmt_table(table)
                    if tt: parts.append(f"[第{i+1}页-表格{ti+1}]\n{tt}")
        content = "\n\n".join(parts)
        if len(content.strip()) < 100 and r["pages"] > 0:
            r["metadata"]["needs_ocr"] = True
        r["content"] = content; r["read_method"] = "pdfplumber"

    def _read_docx(self, fp, r):
        from docx import Document
        doc = Document(fp); parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                if p.style.name.startswith("Heading"):
                    parts.append(f"## {p.text.strip()}")
                else: parts.append(p.text.strip())
        for ti, table in enumerate(doc.tables):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows: parts.append(f"[表格{ti+1}]\n{self._fmt_table(rows)}")
        r["content"] = "\n\n".join(parts)
        r["metadata"] = {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}
        r["pages"] = max(1, len(doc.paragraphs)//30); r["read_method"] = "python-docx"

    def _read_xlsx(self, fp, r):
        from openpyxl import load_workbook
        wb = load_workbook(fp, read_only=True, data_only=True); parts = []
        for sn in wb.sheetnames:
            ws = wb[sn]; rows = []
            for row in ws.iter_rows(values_only=True):
                rd = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in rd): rows.append(rd)
            if rows: parts.append(f"[工作表: {sn}]\n{self._fmt_table(rows)}")
        wb.close()
        r["content"] = "\n\n".join(parts)
        r["metadata"] = {"sheet_count": len(wb.sheetnames)}
        r["pages"] = len(wb.sheetnames); r["read_method"] = "openpyxl"

    def _read_csv(self, fp, r):
        import csv
        with open(fp,"rb") as f:
            enc = chardet.detect(f.read(10000)).get("encoding","utf-8")
        rows = []
        with open(fp,"r",encoding=enc,errors="replace") as f:
            for row in csv.reader(f):
                if any(c.strip() for c in row): rows.append(row)
        r["content"] = self._fmt_table(rows) if rows else ""
        r["metadata"] = {"rows": len(rows), "encoding": enc}
        r["pages"] = 1; r["read_method"] = "csv"

    def _read_text(self, fp, r):
        with open(fp,"rb") as f:
            enc = chardet.detect(f.read(10000)).get("encoding","utf-8")
        with open(fp,"r",encoding=enc,errors="replace") as f:
            r["content"] = f.read()
        r["metadata"] = {"encoding": enc}; r["pages"] = 1; r["read_method"] = "text"

    def _read_image(self, fp, r):
        from PIL import Image
        img = Image.open(fp)
        r["content"] = "[图片文件,需要OCR处理]"
        r["metadata"] = {"width":img.size[0],"height":img.size[1],"format":img.format,"needs_ocr":True}
        r["pages"] = 1; r["read_method"] = "pillow"

    def _fmt_table(self, rows):
        if not rows: return ""
        mc = max(len(r) for r in rows)
        return "\n".join(" | ".join(str(c).strip() for c in (list(r)+[""]*(mc-len(r)))) for r in rows)

    def scan_directory(self, dp):
        files = []
        d = Path(dp)
        if not d.exists(): return files
        for f in sorted(d.iterdir()):
            if f.is_file() and not f.name.startswith("."):
                ext = f.suffix.lower()
                if ext in self.SUPPORTED_TYPES:
                    files.append({"path":str(f),"name":f.name,"type":self.SUPPORTED_TYPES[ext],
                        "size_mb":round(f.stat().st_size/(1024*1024),2),
                        "modified":datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M")})
        return files
